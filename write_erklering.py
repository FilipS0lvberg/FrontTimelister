from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn
import io

def write_erklering1(selskapsnavn, dato, projectnumber, prosjektleder):
    """
    Writes selskapsnavn and dato into the header and footer (right-aligned, bold, Calibri 12pt),
    and writes a project statement in the body using the projectnumber.

    Parameters:
    selskapsnavn (str): Company name to write.
    dato (str): Date to write.
    projectnumber (str): Project number to include in body text.
    filename (str): The name of the file to save (default is 'output.docx').

    Returns:
    str: The path to the saved Word document.
    """
    # Create a new Word document
    doc = Document()

    # Set default font to Calibri size 11 pt for the body
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    # Access the section to modify header and footer
    section = doc.sections[0]
    header = section.header
    footer = section.footer

    # --- Clear existing header and footer paragraphs ---
    for para in header.paragraphs:
        p = para._element
        p.getparent().remove(p)

    for para in footer.paragraphs:
        p = para._element
        p.getparent().remove(p)

    # --- Add header content ---
    p1 = header.add_paragraph()
    run1 = p1.add_run(selskapsnavn)
    run1.bold = True
    run1.font.name = 'Calibri'
    run1.font.size = Pt(12)
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # --- Add footer content (same style) ---
    p4 = footer.add_paragraph()
    run4 = p4.add_run(dato)
    run4.bold = True
    run4.font.name = 'Calibri'
    run4.font.size = Pt(12)
    p4.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # --- Add project statement to body ---
    body_paragraph = doc.add_paragraph()
    body_run = body_paragraph.add_run(
        f""" 

Erklæring prosjektkostnader SkatteFUNN prosjektnummer {projectnumber}."""
    )
    body_run.bold = True
    body_run.font.name = 'Calibri'
    body_run.font.size = Pt(12)

    # --- Add body content statement ---
    body_main_paragraph = doc.add_paragraph()
    body_main_run = body_main_paragraph.add_run(
        """Vi bekrefter at alle rapporterte kostnader/timer i prosjektregnskapet bare relaterer seg til aktiviteter som er godkjent av Forskningsrådet for ovennevnte prosjekt (jamfør søknad og godkjenning) og som utgjør en reell merkostnad for bedriften relatert til prosjektet.

Vi bekrefter at alle bilag som inngår i prosjektregnskapet er regnskapsført løpende og holdt kontroll på slik at det endelige prosjektregnskap ikke inneholder noen bilag/kostnader som danner et uriktig kostnadsgrunnlag for prosjektet.

I tillegg bekrefter vi at timeregistreringen for ovennevnte skattefunnprosjekt er korrekte, herunder at alle oppførte ansatte jobbet i selskapet på tidspunktet, er oppført med korrekt lønn, at samtlige oppførte timer er korrekte og kun relaterer seg til godkjente aktiviteter. Vi bekrefter at oppførte timer er kontrollert med interne rutiner, for å sikre at det totale kostnadsgrunnlaget ikke gir et feilaktig bilde.
"""
    )

    # --- Add spacing before signature ---
    for _ in range(3):  # Add blank lines for spacing
        doc.add_paragraph()

    # --- Add underlined non-breaking spaces for short line ---
    line_paragraph = doc.add_paragraph()
    line_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    non_breaking_spaces = "\u00A0" * 100  # Adjust count for line length
    line_run = line_paragraph.add_run(non_breaking_spaces)
    line_run.underline = True
    line_run.font.name = 'Calibri'
    line_run.font.size = Pt(12)

    # --- Add signature info ---
    signature_paragraph = doc.add_paragraph()
    signature_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    signature_run = signature_paragraph.add_run(
        f"Navn: {prosjektleder}\nStilling: Prosjektleder"
    )
    signature_run.font.name = 'Calibri'
    signature_run.font.size = Pt(12)

    # Save the document
    #doc.save("testing.docx")
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# Example usage
if __name__ == "__main__":
    write_erklering1("Front Innovation AS", "01.01.2025", "123456", "FilipSH")
